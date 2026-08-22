#!/usr/bin/env python3
"""Export the 150-PDF and 22-table RAG evaluations to two Word documents."""
from __future__ import annotations

import json
import re
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports" / "word"
QUESTION_DOCX = OUT_DIR / "RAG_검증_질문목록_172문항.docx"
RESULT_DOCX = OUT_DIR / "RAG_검증_상세결과_172문항.docx"

TEXT_GOLD = ROOT / "data" / "eval" / "broad_pdf_150_final.jsonl"
TEXT_RESULTS = (
    ROOT
    / "data"
    / "processed"
    / "logs"
    / "broad_pdf_150"
    / "final_accurate_v5_judge"
    / "judged_records.jsonl"
)
TABLE_RESULTS = (
    ROOT
    / "data"
    / "processed"
    / "logs"
    / "table_broad_22"
    / "final_fast_v5"
    / "records.jsonl"
)

INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "246B46"
GREEN_FILL = "E9F5EE"
RED = "9B1C1C"
RED_FILL = "FDECEC"
GOLD = "7A5A00"
GOLD_FILL = "FFF6DA"
WHITE = "FFFFFF"
BLACK = "111111"


def read_jsonl(path: Path) -> list[dict]:
    return [
        json.loads(line)
        for line in path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]


def clean_text(value: object) -> str:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    east_asia: str = "Malgun Gothic",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, *, color: str, side: str = "left", size: int = 16) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    border = pbdr.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        pbdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "8")
    border.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:type"), "dxa")
    tblw.set(qn("w:w"), str(total))
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:type"), "dxa")
    tblind.set(qn("w:w"), str(indent_dxa))

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:type"), "dxa")
            tcw.set(qn("w:w"), str(width))
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tcmar = tcpr.find(qn("w:tcMar"))
            if tcmar is None:
                tcmar = OxmlElement("w:tcMar")
                tcpr.append(tcmar)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = tcmar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tcmar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for node in (fld_begin, instr, fld_sep, text, fld_end):
        run._r.append(node)
    set_run_font(run, size=9, color=MUTED)


def apply_style_font(style, *, size: float, color: str, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Malgun Gothic")


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    apply_style_font(normal, size=11, color=BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = doc.styles["Heading 1"]
    apply_style_font(h1, size=16, color=BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    apply_style_font(h2, size=13, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    apply_style_font(h3, size=12, color=DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    specs = {
        "Question List": (11, BLACK, False, 4, 1.25),
        "Item Heading": (11.5, DARK_BLUE, True, 6, 1.15),
        "Field Label": (9, BLUE, True, 2, 1.0),
        "Field Body": (10.5, BLACK, False, 6, 1.18),
        "Answer Section": (9.5, DARK_BLUE, True, 2, 1.1),
        "Answer Body": (9.5, BLACK, False, 3, 1.12),
        "Note Body": (9.5, MUTED, False, 5, 1.12),
        "Table Text": (9.5, BLACK, False, 0, 1.05),
    }
    for name, (size, color, bold, after, spacing) in specs.items():
        if name in doc.styles:
            style = doc.styles[name]
        else:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        apply_style_font(style, size=size, color=color, bold=bold)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = spacing

    doc.styles["Question List"].paragraph_format.left_indent = Inches(0.375)
    doc.styles["Question List"].paragraph_format.first_line_indent = Inches(-0.188)
    doc.styles["Item Heading"].paragraph_format.keep_with_next = True
    doc.styles["Field Label"].paragraph_format.keep_with_next = True
    doc.styles["Answer Body"].paragraph_format.left_indent = Inches(0.18)
    doc.styles["Answer Body"].paragraph_format.right_indent = Inches(0.08)
    doc.styles["Note Body"].paragraph_format.left_indent = Inches(0.18)


def add_numbering(doc: Document, *, bullet: bool = False) -> int:
    root = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in root.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in root.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    level.append(ppr)
    abstract.append(level)
    root.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    root.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)


def add_markdown_runs(paragraph, text: str, *, size: float | None = None) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if bold else part
        content = content.replace("`", "")
        run = paragraph.add_run(content)
        set_run_font(run, size=size, bold=bold)


def add_title_block(doc: Document, title: str, subtitle: str, metadata: str) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(6)
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("MARITIMEOPS AI · VALIDATION RECORD")
    set_run_font(run, size=9, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(title)
    set_run_font(run, size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(subtitle)
    set_run_font(run, size=12.5, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(metadata)
    set_run_font(run, size=9.5, color=MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    set_paragraph_border(rule, color=BLUE, side="bottom", size=10)


def add_header_footer(doc: Document, label: str) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = True
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run_font(run, size=8.5, color=MUTED, bold=True)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("MaritimeOps AI · ")
        set_run_font(run, size=8.5, color=MUTED)
        add_page_field(p)


def new_document(title: str, subject: str, header_label: str) -> Document:
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, header_label)
    doc.core_properties.title = title
    doc.core_properties.subject = subject
    doc.core_properties.author = "MaritimeOps AI Evaluation"
    doc.core_properties.keywords = "RAG, maritime, validation, evaluation"
    return doc


def add_label(doc: Document, label: str) -> None:
    p = doc.add_paragraph(style="Field Label")
    run = p.add_run(label)
    set_run_font(run, size=9, color=BLUE, bold=True)


def add_field_body(
    doc: Document,
    text: str,
    *,
    fill: str | None = None,
    border: str | None = None,
    italic: bool = False,
) -> None:
    lines = [line.strip() for line in clean_text(text).splitlines() if line.strip()]
    if not lines:
        lines = ["기록 없음"]
    for line in lines:
        p = doc.add_paragraph(style="Field Body")
        p.paragraph_format.left_indent = Inches(0.14 if fill or border else 0)
        p.paragraph_format.right_indent = Inches(0.06 if fill else 0)
        if fill:
            set_paragraph_shading(p, fill)
        if border:
            set_paragraph_border(p, color=border, side="left", size=14)
        add_markdown_runs(p, re.sub(r"^[-•]\s*", "", line), size=10.5)
        if italic:
            for run in p.runs:
                run.italic = True


def add_answer_block(doc: Document, answer: str, bullet_num_id: int) -> None:
    lines = clean_text(answer).splitlines()
    wrote = False
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        wrote = True
        if line.startswith("##"):
            p = doc.add_paragraph(style="Answer Section")
            set_paragraph_shading(p, CALLOUT)
            add_markdown_runs(p, re.sub(r"^#+\s*", "", line), size=9.5)
            continue
        if line.startswith("-"):
            p = doc.add_paragraph(style="Answer Body")
            set_paragraph_shading(p, CALLOUT)
            apply_numbering(p, bullet_num_id)
            add_markdown_runs(p, line[1:].strip(), size=9.5)
            continue
        if line.startswith(">"):
            p = doc.add_paragraph(style="Note Body")
            set_paragraph_shading(p, CALLOUT)
            set_paragraph_border(p, color="C9D2DE", side="left", size=10)
            add_markdown_runs(p, line[1:].strip(), size=9.5)
            for run in p.runs:
                run.italic = True
            continue
        p = doc.add_paragraph(style="Answer Body")
        set_paragraph_shading(p, CALLOUT)
        add_markdown_runs(p, line, size=9.5)
    if not wrote:
        p = doc.add_paragraph("실제 답변 기록 없음", style="Note Body")
        set_paragraph_shading(p, CALLOUT)


def text_status(row: dict) -> str:
    if bool(row.get("judge_pass")):
        return "정답"
    if bool(row.get("judge_contradiction")) or int(row.get("judge_correctness") or 0) <= 1:
        return "실패"
    return "보류"


def table_status(row: dict) -> str:
    if bool(row.get("value_hit")):
        return "정답"
    if not row.get("error") and row.get("evidence_file_hit") and row.get("evidence_page_hit"):
        return "보류"
    return "실패"


def status_colors(status: str) -> tuple[str, str]:
    if status == "정답":
        return GREEN, GREEN_FILL
    if status == "실패":
        return RED, RED_FILL
    return GOLD, GOLD_FILL


def add_status_block(doc: Document, status: str, details: str) -> None:
    color, fill = status_colors(status)
    p = doc.add_paragraph(style="Field Body")
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.06)
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, color=color, side="left", size=18)
    run = p.add_run(f"{status}  ")
    set_run_font(run, size=10.5, color=color, bold=True)
    run = p.add_run(clean_text(details))
    set_run_font(run, size=9.5, color=BLACK)


def add_item_heading(doc: Document, title: str, status: str) -> None:
    color, fill = status_colors(status)
    p = doc.add_paragraph(style="Item Heading")
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, color=color, side="left", size=18)
    run = p.add_run(title)
    set_run_font(run, size=11.5, color=DARK_BLUE, bold=True)
    run = p.add_run(f"  |  {status}")
    set_run_font(run, size=11.5, color=color, bold=True)


def add_summary_table(doc: Document, rows: list[tuple[str, int, int, int, int]]) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["구분", "정답", "실패", "보류", "합계"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = "Table Text"
        run = p.add_run(value)
        set_run_font(run, size=9.5, color=INK, bold=True)
    set_repeat_table_header(table.rows[0])
    for label, correct, failed, hold, total in rows:
        cells = table.add_row().cells
        for index, value in enumerate((label, correct, failed, hold, total)):
            p = cells[index].paragraphs[0]
            p.style = "Table Text"
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_font(run, size=9.5, color=BLACK, bold=index == 4)
    set_table_geometry(table, [2500, 1715, 1715, 1715, 1715])
    doc.add_paragraph()


def build_question_document(text_gold: list[dict], table_rows: list[dict]) -> Document:
    doc = new_document(
        "RAG 검증 질문 목록 172문항",
        "150개 PDF 텍스트 질문과 22개 표 질문",
        "RAG VALIDATION · QUESTION LIST",
    )
    add_title_block(
        doc,
        "RAG 검증 질문 목록",
        "서로 다른 PDF 150문항 + 표 질문 22문항",
        f"총 172문항 · 작성일 {date.today().isoformat()} · 답변과 판정은 제외",
    )
    p = doc.add_paragraph(style="Note Body")
    set_paragraph_shading(p, CALLOUT)
    set_paragraph_border(p, color=BLUE, side="left", size=14)
    run = p.add_run("이 문서는 UI 사용자 검증용 질문만 수록합니다. 질문 ID는 상세 결과 문서와 동일합니다.")
    set_run_font(run, size=9.5, color=MUTED)

    text_num = add_numbering(doc)
    doc.add_heading("1. 텍스트 PDF 질문 (150문항)", level=1)
    by_source: dict[str, list[dict]] = defaultdict(list)
    for row in text_gold:
        by_source[str(row.get("gold_source") or "기타")].append(row)
    source_order = ["DNV", "MEPC", "MSC", "KR", "ABS", "LR"]
    for source in source_order:
        rows = by_source.get(source, [])
        if not rows:
            continue
        doc.add_heading(f"{source} ({len(rows)}문항)", level=2)
        for row in rows:
            p = doc.add_paragraph(style="Question List")
            apply_numbering(p, text_num)
            run = p.add_run(f"[{row['question_id']}] ")
            set_run_font(run, size=10.5, color=BLUE, bold=True)
            run = p.add_run(clean_text(row.get("question")))
            set_run_font(run, size=11, color=BLACK)

    table_num = add_numbering(doc)
    doc.add_heading("2. 표 질문 (22문항)", level=1)
    for row in table_rows:
        p = doc.add_paragraph(style="Question List")
        apply_numbering(p, table_num)
        run = p.add_run(f"[{row.get('qid')}] ")
        set_run_font(run, size=10.5, color=BLUE, bold=True)
        run = p.add_run(clean_text(row.get("question")))
        set_run_font(run, size=11, color=BLACK)
    return doc


def build_result_document(
    text_gold: list[dict], text_results: list[dict], table_rows: list[dict]
) -> Document:
    doc = new_document(
        "RAG 검증 상세 결과 172문항",
        "기준 정답, 실제 답변, 자동 판정과 근거",
        "RAG VALIDATION · DETAILED RESULTS",
    )
    add_title_block(
        doc,
        "RAG 검증 상세 결과",
        "질문 · 기준 정답 · 실제 생성 답변 · 정답/실패/보류",
        "텍스트: Accurate final_accurate_v5 · 표: Fast final_fast_v5 · 자동 판정 기록",
    )

    p = doc.add_paragraph(style="Note Body")
    set_paragraph_shading(p, GOLD_FILL)
    set_paragraph_border(p, color=GOLD, side="left", size=16)
    run = p.add_run(
        "판정 기준: 정답은 자동 근거 판정 통과, 실패는 정확성 0~1점 또는 명시적 모순, "
        "보류는 부분 충족·표기 차이 등 사람 확인이 필요한 경우입니다. 해사 전문가의 공식 검수 결과는 아닙니다."
    )
    set_run_font(run, size=9.5, color=BLACK)

    result_by_id = {row["question_id"]: row for row in text_results}
    merged_text: list[tuple[dict, dict, str]] = []
    for gold in text_gold:
        result = result_by_id[gold["question_id"]]
        merged_text.append((gold, result, text_status(result)))

    text_counts = Counter(status for _, _, status in merged_text)
    table_counts = Counter(table_status(row) for row in table_rows)
    total_counts = text_counts + table_counts
    add_summary_table(
        doc,
        [
            ("텍스트 150", text_counts["정답"], text_counts["실패"], text_counts["보류"], 150),
            ("표 22", table_counts["정답"], table_counts["실패"], table_counts["보류"], 22),
            ("전체", total_counts["정답"], total_counts["실패"], total_counts["보류"], 172),
        ],
    )

    bullet_num_id = add_numbering(doc, bullet=True)
    doc.add_heading("1. 텍스트 PDF 질문 상세 결과", level=1)
    by_source: dict[str, list[tuple[dict, dict, str]]] = defaultdict(list)
    for entry in merged_text:
        by_source[str(entry[0].get("gold_source") or "기타")].append(entry)

    source_order = ["DNV", "MEPC", "MSC", "KR", "ABS", "LR"]
    for source in source_order:
        entries = by_source.get(source, [])
        if not entries:
            continue
        counts = Counter(status for _, _, status in entries)
        doc.add_heading(
            f"{source} · {len(entries)}문항 (정답 {counts['정답']} / 실패 {counts['실패']} / 보류 {counts['보류']})",
            level=2,
        )
        for gold, result, status in entries:
            add_item_heading(doc, gold["question_id"], status)
            add_label(doc, "질문")
            add_field_body(doc, gold.get("question"), fill=WHITE, border=BLUE)
            add_label(doc, "기준 정답")
            add_field_body(doc, gold.get("gold_answer"), fill=LIGHT_BLUE, border=BLUE)
            add_label(doc, "실제 Accurate 답변")
            add_answer_block(doc, result.get("answer"), bullet_num_id)
            add_label(doc, "자동 판정")
            score_text = (
                f"정확성 {result.get('judge_correctness', 0)}/4 · "
                f"완전성 {result.get('judge_completeness', 0)}/4 · "
                f"직접응답 {'예' if result.get('judge_relevance') else '아니오'} · "
                f"모순 {'있음' if result.get('judge_contradiction') else '없음'} · "
                f"근거 없는 구체 주장 {'있음' if result.get('judge_unsupported_specific_claim') else '없음'}"
            )
            reason = clean_text(result.get("judge_reason"))
            details = score_text + (f" · {reason}" if reason else "")
            add_status_block(doc, status, details)
            add_label(doc, "정답 근거")
            pages = ", ".join(str(page) for page in gold.get("gold_pages") or []) or str(
                gold.get("gold_page") or "-"
            )
            add_field_body(
                doc,
                f"{gold.get('gold_file_name') or '-'} · p.{pages}",
                fill=LIGHT_GRAY,
            )

    doc.add_heading("2. 표 질문 상세 결과", level=1)
    for row in table_rows:
        status = table_status(row)
        add_item_heading(doc, str(row.get("qid") or "TABLE"), status)
        add_label(doc, "질문")
        add_field_body(doc, row.get("question"), fill=WHITE, border=BLUE)
        add_label(doc, "기준 정답")
        add_field_body(doc, row.get("gold_answer"), fill=LIGHT_BLUE, border=BLUE)
        add_label(doc, "실제 Fast 답변")
        add_answer_block(doc, row.get("answer"), bullet_num_id)
        add_label(doc, "자동 판정")
        if status == "정답":
            details = "정답값이 일치했습니다."
        elif status == "보류":
            details = (
                "엄격 문자열 비교에서는 불일치했지만 정답 문서와 페이지는 적중했습니다. "
                f"기준값 ‘{clean_text(row.get('gold_answer'))}’과 실제 표기 차이를 사람이 확인해야 합니다."
            )
        else:
            details = clean_text(row.get("error")) or "정답값 또는 근거 문서·페이지가 일치하지 않았습니다."
        details += (
            f" · 문서 적중 {'예' if row.get('evidence_file_hit') else '아니오'}"
            f" · 페이지 적중 {'예' if row.get('evidence_page_hit') else '아니오'}"
        )
        add_status_block(doc, status, details)
        add_label(doc, "정답 근거")
        add_field_body(
            doc,
            f"{row.get('gold_file_name') or '-'} · p.{row.get('gold_page') or '-'}",
            fill=LIGHT_GRAY,
        )
    return doc


def audit_document(doc: Document, *, expected_questions: int | None = None) -> None:
    section = doc.sections[0]
    assert abs(section.page_width.inches - 8.5) < 0.01
    assert abs(section.page_height.inches - 11) < 0.01
    assert abs(section.left_margin.inches - 1.0) < 0.01
    assert abs(section.right_margin.inches - 1.0) < 0.01
    for table in doc.tables:
        tblpr = table._tbl.tblPr
        tblw = tblpr.find(qn("w:tblW"))
        assert tblw is not None and tblw.get(qn("w:type")) == "dxa"
        assert sum(int(col.get(qn("w:w"))) for col in table._tbl.tblGrid) == int(
            tblw.get(qn("w:w"))
        )
    if expected_questions is not None:
        text = "\n".join(p.text for p in doc.paragraphs)
        assert text.count("[BPDF-") == 150
        assert text.count("[TC22_") == 22


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    text_gold = read_jsonl(TEXT_GOLD)
    text_results = read_jsonl(TEXT_RESULTS)
    table_rows = read_jsonl(TABLE_RESULTS)
    assert len(text_gold) == 150
    assert len(text_results) == 150
    assert len(table_rows) == 22
    assert {row["question_id"] for row in text_gold} == {
        row["question_id"] for row in text_results
    }

    questions = build_question_document(text_gold, table_rows)
    audit_document(questions, expected_questions=172)
    questions.save(QUESTION_DOCX)

    results = build_result_document(text_gold, text_results, table_rows)
    audit_document(results)
    results.save(RESULT_DOCX)

    counts = Counter(text_status(row) for row in text_results)
    counts.update(table_status(row) for row in table_rows)
    print(
        json.dumps(
            {
                "question_docx": str(QUESTION_DOCX),
                "result_docx": str(RESULT_DOCX),
                "text_questions": len(text_gold),
                "table_questions": len(table_rows),
                "status_counts": dict(counts),
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
