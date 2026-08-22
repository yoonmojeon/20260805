#!/usr/bin/env python3
"""Export a Word handout of questions the pipeline already answers correctly.

The default output is meant to be handed to people who will type the questions
into the UI, so it carries questions only — no gold answers, no scores, no
timings. Use --with-answers for the presenter's copy.

Selection is evidence-based: only questions that passed in the stored evaluation
runs are used, never hand-picked guesses.

- 표: data/eval/quality_50_open_mix_gemma4_12b_current.json,
      data/eval/balanced_quality_100_gemma4_12b_search_final.json
      (needle PASS + quality GOOD + route_ok)
- 텍스트: data/processed/logs/text_rag_eval_v3/gemma4_12b_full405_current_20260815
      (quality_score 1.0 + completeness 1.0 + behavior_pass + 금지주장 없음)

Questions whose expected behaviour is a refusal, or whose wording leaks internal
jargon such as 청크, are excluded: a live demo should produce real answers.

Layout is a numbered question block per item. Wide multi-column tables were
dropped because Word autofit ignores per-cell widths and squeezes long Korean
questions into unreadable columns.
"""
from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
TABLE_SET = ROOT / "data" / "eval" / "table_questions_22docs_practical_v1_curated.jsonl"
TEXT_SET = ROOT / "data" / "eval" / "pilot_validation_text_v3.jsonl"
TABLE_RESULTS = [
    ROOT / "data" / "eval" / "quality_50_open_mix_gemma4_12b_current.json",
    ROOT / "data" / "eval" / "balanced_quality_100_gemma4_12b_search_final.json",
]
TEXT_RECORDS = (
    ROOT
    / "data"
    / "processed"
    / "logs"
    / "text_rag_eval_v3"
    / "gemma4_12b_full405_current_20260815"
    / "records.jsonl"
)
VERIFY_DIR = ROOT / "data" / "processed" / "logs" / "demo_question_verify"
TABLE_VERDICTS = VERIFY_DIR / "table_verdicts.json"

KO_FONT = "맑은 고딕"
GRAY = RGBColor(0x44, 0x44, 0x44)
BLUE = RGBColor(0x1F, 0x4E, 0x79)

TEST_TYPE_KO = {
    "seed": "시드",
    "paraphrase": "표현 변형",
    "evidence_precision": "근거 정밀",
    "noise_robustness": "노이즈 내성",
    "negative_rejection": "근거 없음 거절",
    "counterfactual_robustness": "잘못된 전제 교정",
    "boundary": "유사 문서 구분",
    "format": "출력 형식",
    "scope": "적용 범위",
    "integration": "다문서 통합",
}
# Answering ability first, refusal-style checks last: a demo should lead with answers.
TEXT_TYPE_ORDER = [
    "seed",
    "evidence_precision",
    "paraphrase",
    "counterfactual_robustness",
    "noise_robustness",
    "format",
    "scope",
    "negative_rejection",
    "boundary",
    "integration",
]
SCENARIO_KO = {
    "V01": "최신 동향 / MEPC 84-7-14",
    "V02": "회의 결과 / MSC 111-WP.1",
    "V03": "환경규제 / MEPC 84-6-2",
    "V04": "환경규제 / MSC 111 대체연료",
    "V05": "자율운항 / MSC 111-5",
    "V06": "선급 Rule / DNV-CG-0264",
    "V07": "선급 Rule / LR Notice No.1",
    "V08": "선급 Rule / ABS Smart Functions",
    "V09": "선급 Rule / ABS Autonomous",
}


def read_jsonl(path: Path) -> list[dict]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def norm(text: str) -> str:
    return " ".join(str(text or "").split())


def shorten(text: str, limit: int) -> str:
    flat = norm(text)
    return flat if len(flat) <= limit else flat[: limit - 1] + "…"


JARGON = re.compile(r"청크|검색 후보|섞여도|hard[- ]?negative", re.I)

def fix_particles(text: str) -> str:
    """Correct the object particle 을/를 against the preceding syllable's final consonant.

    The generated eval sets slot document titles into templates, so strings like
    '보고서 초안를' and '원시데이터을' appear. Only 을/를 is corrected: 은/는 and 이/가
    also occur as verb endings ('접하지 않는', '있는'), where rewriting them is wrong.
    """
    chars = list(text)
    for i, ch in enumerate(chars):
        if ch not in ("을", "를") or i == 0:
            continue
        if i + 1 < len(chars) and chars[i + 1] not in " \t":
            continue
        prev = chars[i - 1]
        if not ("가" <= prev <= "힣"):
            continue
        has_final = (ord(prev) - 0xAC00) % 28 != 0
        chars[i] = "을" if has_final else "를"
    return "".join(chars)


def load_verified_text_answers() -> dict[str, str]:
    """Answers the model actually produced in the latest re-verification.

    The presenter's copy shows these instead of the gold wording, because the
    automatic completeness check accepts alias matches and the gold phrasing can
    differ from what the model says.
    """
    dirs = sorted(VERIFY_DIR.glob("text10_*/records.jsonl"))
    if not dirs:
        return {}
    return {rec["question_id"]: rec.get("answer") or "" for rec in read_jsonl(dirs[-1])}


def answer_gist(answer: str, limit: int = 240) -> str:
    """First substantive bullets of a sectioned answer, flattened to one line."""
    bullets = [
        line.strip().lstrip("-*").strip()
        for line in str(answer or "").splitlines()
        if line.strip().startswith("-")
    ]
    body = " / ".join(b for b in bullets if b)[:600] or norm(answer)
    return shorten(body.replace("**", ""), limit)


def load_table_verdicts() -> tuple[set[str], dict[str, float]]:
    """Regressed questions and measured seconds from the latest re-verification."""
    if not TABLE_VERDICTS.exists():
        return set(), {}
    data = json.loads(TABLE_VERDICTS.read_text(encoding="utf-8"))
    failed = {norm(q) for q in data.get("failed", [])}
    seconds = {norm(q): float(v) for q, v in (data.get("seconds") or {}).items()}
    return failed, seconds


def pick_table_questions(want: int, max_per_doc: int = 3, only_qids: list[str] | None = None) -> list[dict]:
    """Passing table questions, fastest answers first, spread over source documents.

    only_qids pins a hand-reviewed selection and keeps its order.
    """
    curated = {norm(row["question"]): row for row in read_jsonl(TABLE_SET)}
    regressed, verified_seconds = load_table_verdicts()

    passing: list[dict] = []
    seen: set[str] = set()
    for path in TABLE_RESULTS:
        if not path.exists():
            continue
        for res in json.loads(path.read_text(encoding="utf-8"))["results"]:
            if not str(res.get("type", "")).startswith("table"):
                continue
            if res.get("needle") != "PASS" or res.get("quality") != "GOOD" or not res.get("route_ok"):
                continue
            key = norm(res["question"])
            if key in seen or key in regressed:
                continue
            # Debug-style questions spell out the file name and page; a demo needs
            # questions a real user would type.
            if re.search(r"\.pdf|\d+\s*(?:쪽|페이지)", key):
                continue
            seen.add(key)
            passing.append({"result": res, "gold_row": curated.get(key), "run": path.stem})

    if only_qids:
        by_qid = {(item["gold_row"] or {}).get("qid"): item for item in passing}
        return [by_qid[qid] for qid in only_qids if qid in by_qid]

    passing.sort(
        key=lambda item: verified_seconds.get(
            norm(item["result"]["question"]), float(item["result"].get("dt") or 999)
        )
    )
    picked: list[dict] = []
    doc_usage: Counter[str] = Counter()
    for item in passing:
        doc = (item["gold_row"] or {}).get("gold_file_name") or ""
        if doc and doc_usage[doc] >= max_per_doc:
            continue
        doc_usage[doc] += 1
        picked.append(item)
        if len(picked) >= want:
            break
    for item in passing:  # relax the per-document cap only if still short
        if len(picked) >= want:
            break
        if item not in picked:
            picked.append(item)
    return picked[:want]


def trigrams(text: str) -> set[str]:
    flat = norm(text).replace(" ", "")
    return {flat[i : i + 3] for i in range(max(len(flat) - 2, 1))}


def too_similar(candidate: str, chosen: list[str], limit: float = 0.6) -> bool:
    cand = trigrams(candidate)
    for other in chosen:
        ref = trigrams(other)
        union = cand | ref
        if union and len(cand & ref) / len(union) >= limit:
            return True
    return False


def pick_text_questions(want: int, only_ids: list[str] | None = None) -> list[dict]:
    """Perfect-score text questions, balanced over scenarios and test types.

    Balancing both axes at once matters: negative_rejection alone has dozens of
    perfect records, and a few scenarios dominate, so an unweighted pick repeats
    the same document and the same wording. only_ids pins a hand-reviewed
    selection and keeps its order.
    """
    eval_rows = {row["question_id"]: row for row in read_jsonl(TEXT_SET)}

    pool = [
        rec
        for rec in read_jsonl(TEXT_RECORDS)
        if not rec.get("error")
        and rec.get("behavior_pass")
        and not rec.get("forbidden_violations")
        and rec.get("quality_score") == 1.0
        and rec.get("completeness") == 1.0
        and rec.get("test_type") != "negative_rejection"
        and not JARGON.search(rec.get("question") or "")
    ]
    type_rank = {name: idx for idx, name in enumerate(TEXT_TYPE_ORDER)}

    if only_ids:
        by_id = {rec["question_id"]: rec for rec in pool}
        return [
            {"record": by_id[qid], "gold_row": eval_rows.get(qid)} for qid in only_ids if qid in by_id
        ]

    picked: list[dict] = []
    chosen_questions: list[str] = []
    scenario_usage: Counter[str] = Counter()
    type_usage: Counter[str] = Counter()
    while pool and len(picked) < want:
        pool.sort(
            key=lambda r: (
                scenario_usage[str(r.get("scenario_id"))],
                type_usage[str(r.get("test_type"))],
                float(r.get("e2e_seconds") or 999) >= 15,  # keep a live demo snappy
                type_rank.get(str(r.get("test_type")), 99),
                float(r.get("e2e_seconds") or 999),
                str(r.get("question_id")),
            )
        )
        rec = next(
            (r for r in pool if not too_similar(r.get("question") or "", chosen_questions)),
            None,
        )
        if rec is None:
            break
        pool.remove(rec)
        scenario_usage[str(rec.get("scenario_id"))] += 1
        type_usage[str(rec.get("test_type"))] += 1
        chosen_questions.append(rec.get("question") or "")
        picked.append({"record": rec, "gold_row": eval_rows.get(rec["question_id"])})
    picked.sort(key=lambda item: type_rank.get(str(item["record"].get("test_type")), 99))
    return picked


def text_expectation(rec: dict, gold_row: dict | None) -> str:
    row = gold_row or {}
    if str(rec.get("expected_behavior")) == "refuse_no_evidence" or row.get("answerability") is False:
        return "근거 없음을 밝히고 추측하지 않아야 함"
    if row.get("false_premise"):
        return f"전제 오류를 교정해야 함 — {shorten(row['false_premise'], 110)}"
    points = row.get("gold_answer_points") or []
    if points:
        return " / ".join(shorten(p.get("text", ""), 110) for p in points[:2])
    return shorten(row.get("gold_answer") or "", 220)


def apply_ko_font(run) -> None:
    run.font.name = KO_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KO_FONT)


def setup_styles(document: Document) -> None:
    for name, size in (("Normal", 10), ("Title", 20), ("Heading 1", 14), ("Heading 2", 11)):
        style = document.styles[name]
        style.font.name = KO_FONT
        style.font.size = Pt(size)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), KO_FONT)
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    for attr in ("left_margin", "right_margin"):
        setattr(section, attr, Cm(2))
    for attr in ("top_margin", "bottom_margin"):
        setattr(section, attr, Cm(1.8))


def add_para(
    document: Document,
    text: str,
    *,
    size: float = 10,
    bold: bool = False,
    color: RGBColor | None = None,
    indent: float = 0.0,
    space_after: float = 2,
    align=None,
) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    apply_ko_font(run)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(indent)
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(0)
    if align is not None:
        fmt.alignment = align


def add_heading(document: Document, text: str, level: int) -> None:
    heading = document.add_heading(level=level)
    apply_ko_font(heading.add_run(text))
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(6)


def build_table_blocks(document: Document, items: list[dict], md: list[str], with_answers: bool) -> None:
    for idx, item in enumerate(items, start=1):
        res = item["result"]
        gold_row = item["gold_row"] or {}
        question = fix_particles(shorten(res.get("question") or "", 300))
        add_para(document, f"{idx}. {question}", size=11, space_after=2 if with_answers else 9)
        md.append(f"{idx}. {question}")
        if not with_answers:
            continue
        answer = shorten(res.get("gold") or "", 120)
        evidence = f"{gold_row.get('gold_file_name')} p.{gold_row.get('gold_page')}"
        caption = shorten(gold_row.get("table_caption") or "", 70)
        if caption:
            evidence += f" · {caption}"
        add_para(document, f"정답 · {answer}", size=9, color=GRAY, indent=0.7, space_after=1)
        add_para(document, f"근거 · {evidence}", size=9, color=GRAY, indent=0.7, space_after=9)
        md += [f"    - 정답: {answer}", f"    - 근거: {evidence}"]


def build_text_blocks(document: Document, items: list[dict], md: list[str], with_answers: bool) -> None:
    verified = load_verified_text_answers() if with_answers else {}
    for idx, item in enumerate(items, start=1):
        rec = item["record"]
        question = fix_particles(shorten(rec.get("question") or "", 300))
        add_para(document, f"{idx}. {question}", size=11, space_after=2 if with_answers else 9)
        md.append(f"{idx}. {question}")
        if not with_answers:
            continue
        answer = verified.get(str(rec.get("question_id")))
        label = "실제 답변" if answer else "기대"
        expectation = answer_gist(answer) if answer else fix_particles(text_expectation(rec, item["gold_row"]))
        scenario = SCENARIO_KO.get(str(rec.get("scenario_id")), str(rec.get("scenario_id")))
        type_ko = TEST_TYPE_KO.get(str(rec.get("test_type")), str(rec.get("test_type")))
        add_para(document, f"{label} · {expectation}", size=9, color=GRAY, indent=0.7, space_after=1)
        add_para(document, f"문서 · {scenario}   |   유형 · {type_ko}", size=9, color=GRAY, indent=0.7, space_after=9)
        md += [f"    - {label}: {expectation}", f"    - 문서: {scenario} | 유형: {type_ko}"]


def save_docx(document: Document, target: Path) -> Path:
    """Save, falling back to a numbered name while the previous file is open in a viewer."""
    for attempt in range(1, 10):
        path = target if attempt == 1 else target.with_name(f"{target.stem}_{attempt}{target.suffix}")
        try:
            document.save(path)
            return path
        except PermissionError:
            continue
    raise PermissionError(f"cannot write {target.name}; close the open copy and retry")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--count", type=int, default=10, help="문항 수 (표·텍스트 각각)")
    parser.add_argument("--with-answers", action="store_true", help="발표자용 정답·근거 포함본")
    parser.add_argument(
        "--picks",
        type=Path,
        help="답변 원문을 검토해 고정한 선정 파일 (예: data/eval/demo_top4_picks.json)",
    )
    args = parser.parse_args()
    with_answers = args.with_answers

    picks = json.loads(args.picks.read_text(encoding="utf-8")) if args.picks else {}
    table_qids = picks.get("table_qids") or None
    text_ids = picks.get("text_question_ids") or None
    count = len(table_qids) if table_qids else args.count

    table_items = pick_table_questions(count, only_qids=table_qids)
    text_items = pick_text_questions(count, only_ids=text_ids)

    stem = "질문리스트_시연용" if not picks else f"질문리스트_시연용_{count}개씩"
    if with_answers:
        stem += "_정답지"
    out_docx = ROOT / "docs" / f"{stem}.docx"
    out_md = out_docx.with_suffix(".md")

    document = Document()
    setup_styles(document)

    title = document.add_heading(level=0)
    apply_ko_font(title.add_run("문서 검색 시연 질문 리스트"))
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        document,
        "아래 질문을 그대로 입력해 보세요.",
        size=10,
        color=GRAY,
        space_after=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    md: list[str] = ["# 문서 검색 시연 질문 리스트", "", "아래 질문을 그대로 입력해 보세요.", ""]

    add_heading(document, "표에서 값을 찾는 질문", 1)
    add_para(document, "선급 규칙·지침의 표에서 특정 값을 찾아옵니다.", size=9.5, color=GRAY, space_after=10)
    md += ["## 표에서 값을 찾는 질문", "", "선급 규칙·지침의 표에서 특정 값을 찾아옵니다.", ""]
    build_table_blocks(document, table_items, md, with_answers)

    add_heading(document, "문서 내용을 묻는 질문", 1)
    add_para(document, "IMO 회의 문서와 선급 Rule 본문을 읽고 정리합니다.", size=9.5, color=GRAY, space_after=10)
    md += ["", "## 문서 내용을 묻는 질문", "", "IMO 회의 문서와 선급 Rule 본문을 읽고 정리합니다.", ""]
    build_text_blocks(document, text_items, md, with_answers)

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    out_docx = save_docx(document, out_docx)
    out_md = out_docx.with_suffix(".md")
    out_md.write_text("\n".join(md) + "\n", encoding="utf-8")
    print(f"wrote {out_docx.name} and {out_md.name} table={len(table_items)} text={len(text_items)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
