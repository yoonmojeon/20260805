"""Build the 50-question UI demo list from verified text and table runs."""
from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = Path(r"C:\Users\user\Desktop\질문리스트.docx")
TABLE_GOLD = ROOT / "data/eval/table_questions_22docs_practical_v1_curated.jsonl"
OUTPUT_DOCX = ROOT / "artifacts/질문리스트.docx"
MANIFEST = ROOT / "artifacts/질문리스트_검증선정.json"

WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

# These positions passed the current Accurate UI path semantically.  Questions
# with weaker source attribution or less representative demo value are omitted
# to make room for table QA without changing the requested total of 50.
TEXT_POSITIONS = (
    1, 2, 3, 5, 6, 8, 9, 10, 11, 12,
    13, 15, 16, 18, 19, 21, 22, 24, 27, 28,
    29, 30, 32, 33, 34, 35, 36, 37, 38, 39,
    40, 41, 43, 44, 45, 46, 47, 48, 49, 50,
)

# Each selected table question passed value + source PDF + source page through
# services.rag_service.run_rag_query(..., latency_mode="accurate").
TABLE_QIDS = (
    "TC22_001",
    "TC22_004",
    "TC22_007",
    "TC22_013",
    "TC22_016",
    "TC22_019",
    "TC22_025",
    "TC22_034",
    "TC22_052",
    "TC22_064",
)


def read_existing_questions(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    questions: list[str] = []
    for paragraph in root.findall(".//w:body/w:p", WORD_NS):
        text = "".join(
            node.text or "" for node in paragraph.findall(".//w:t", WORD_NS)
        ).strip()
        if not text or text == "질문리스트" or text.startswith("텍스트·규정·회의 질문"):
            continue
        if text.startswith("표 기반 질문"):
            continue
        text = re.sub(r"^\s*\d+\s*[.)]\s*", "", text).strip()
        if text:
            questions.append(text)
    return questions


def read_table_rows(path: Path) -> dict[str, dict]:
    rows = [
        json.loads(line)
        for line in path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]
    return {str(row["qid"]): row for row in rows}


def set_cell_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)

    if "Question" not in styles:
        question_style = styles.add_style("Question", WD_STYLE_TYPE.PARAGRAPH)
    else:
        question_style = styles["Question"]
    question_style.base_style = normal
    question_style.font.name = "맑은 고딕"
    question_style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    question_style.font.size = Pt(10.5)
    question_style.paragraph_format.space_after = Pt(5)
    question_style.paragraph_format.line_spacing = 1.14
    question_style.paragraph_format.keep_together = True
    question_style.paragraph_format.widow_control = True


def build_document(text_questions: list[str], table_rows: list[dict]) -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)
    configure_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(11)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("질문리스트")
    run.bold = True
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(15, 75, 88)

    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(2)
    heading.paragraph_format.space_after = Pt(7)
    heading.paragraph_format.keep_with_next = True
    set_cell_shading(heading, "DDEFF2")
    heading_run = heading.add_run("텍스트·규정·회의 질문 (1–40)")
    heading_run.bold = True
    heading_run.font.size = Pt(11.5)
    heading_run.font.color.rgb = RGBColor(15, 75, 88)

    for number, question in enumerate(text_questions, 1):
        paragraph = document.add_paragraph(style="Question")
        paragraph.paragraph_format.left_indent = Cm(0.2)
        paragraph.paragraph_format.first_line_indent = Cm(-0.2)
        number_run = paragraph.add_run(f"{number}. ")
        number_run.bold = True
        paragraph.add_run(question)

    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(7)
    heading.paragraph_format.keep_with_next = True
    set_cell_shading(heading, "DDEFF2")
    heading_run = heading.add_run("표 기반 질문 (41–50)")
    heading_run.bold = True
    heading_run.font.size = Pt(11.5)
    heading_run.font.color.rgb = RGBColor(15, 75, 88)

    for number, row in enumerate(table_rows, 41):
        paragraph = document.add_paragraph(style="Question")
        paragraph.paragraph_format.left_indent = Cm(0.2)
        paragraph.paragraph_format.first_line_indent = Cm(-0.2)
        number_run = paragraph.add_run(f"{number}. ")
        number_run.bold = True
        paragraph.add_run(str(row["question"]).strip())

    add_page_number(section.footer.paragraphs[0])
    document.core_properties.title = "질문리스트"
    document.core_properties.subject = "RAG UI 검증용 질문 50문항"
    document.core_properties.author = ""
    return document


def main() -> None:
    existing = read_existing_questions(SOURCE_DOCX)
    if len(existing) != 50:
        raise RuntimeError(f"기존 질문 수가 50이 아닙니다: {len(existing)}")
    text_questions = [existing[position - 1] for position in TEXT_POSITIONS]
    table_by_qid = read_table_rows(TABLE_GOLD)
    table_rows = [table_by_qid[qid] for qid in TABLE_QIDS]
    if len(text_questions) != 40 or len(table_rows) != 10:
        raise RuntimeError("최종 질문 구성이 40+10이 아닙니다.")

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document = build_document(text_questions, table_rows)
    document.save(OUTPUT_DOCX)
    manifest = {
        "total": 50,
        "text_count": 40,
        "table_count": 10,
        "text_source_positions": list(TEXT_POSITIONS),
        "table_qids": list(TABLE_QIDS),
        "questions": [*text_questions, *(str(row["question"]) for row in table_rows)],
    }
    MANIFEST.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(OUTPUT_DOCX)
    print(MANIFEST)


if __name__ == "__main__":
    main()
