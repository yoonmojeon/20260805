"""
agent/cii.py 단위 테스트 + DWT DB 동기화 검증 + Word 보고서 통합 테스트.

실제 운항 센서 데이터나 프로덕션 DB(data/maritime.db)는 건드리지 않는다.
DB 관련 테스트는 tmp_path 임시 SQLite 파일을 사용한다.
"""
import math
import re
from pathlib import Path

import pytest

from config import CII_PARAMS, VESSEL
from agent import cii


TOL = 1e-9


def approx(a, b, tol=1e-6):
    return math.isclose(a, b, rel_tol=tol, abs_tol=tol)


# ── 1. Bulk Carrier Reference CII ───────────────────────────────────────────
def test_reference_cii_bulk_carrier():
    ref = cii.reference_cii(81200)
    assert approx(ref, 4.192686272987268)
    # 지수는 -c 여야 한다 (a * DWT^(+c) 로 계산하지 않음 검증)
    wrong = CII_PARAMS["a"] * (81200 ** CII_PARAMS["c"])
    assert not approx(ref, wrong)


# ── 2 & 3. Required CII (2025 / 2026) ───────────────────────────────────────
def test_required_cii_2025():
    ref = cii.reference_cii(81200)
    req = cii.required_cii(ref, cii.reduction_factor_percent(2025))
    assert approx(req, 3.815344508418414)


def test_required_cii_2026():
    ref = cii.reference_cii(81200)
    req = cii.required_cii(ref, cii.reduction_factor_percent(2026))
    assert approx(req, 3.731490782958668)


def test_reduction_factors_exact_values():
    expected = {
        2023: 5.000, 2024: 7.000, 2025: 9.000, 2026: 11.000,
        2027: 13.625, 2028: 16.250, 2029: 18.875, 2030: 21.500,
    }
    for year, z in expected.items():
        assert cii.reduction_factor_percent(year) == z


# ── 4. 등급 내부 값 (A~E) ────────────────────────────────────────────────────
@pytest.fixture
def required_2026():
    ref = cii.reference_cii(81200)
    return cii.required_cii(ref, cii.reduction_factor_percent(2026))


def test_rating_inside_a(required_2026):
    attained = required_2026 * 0.80   # < d1(0.86)
    assert cii.rate(attained, required_2026) == "A"


def test_rating_inside_b(required_2026):
    attained = required_2026 * 0.90   # d1~d2 (0.86~0.94)
    assert cii.rate(attained, required_2026) == "B"


def test_rating_inside_c(required_2026):
    attained = required_2026 * 1.00   # d2~d3 (0.94~1.06)
    assert cii.rate(attained, required_2026) == "C"


def test_rating_inside_d(required_2026):
    attained = required_2026 * 1.12   # d3~d4 (1.06~1.18)
    assert cii.rate(attained, required_2026) == "D"


def test_rating_inside_e(required_2026):
    attained = required_2026 * 1.30   # > d4(1.18)
    assert cii.rate(attained, required_2026) == "E"


# ── 5. 정확한 경계값 → 더 좋은 등급 ─────────────────────────────────────────
def test_boundary_d1_is_a(required_2026):
    d1 = CII_PARAMS["d_factors"]["d1"]
    assert cii.rate(required_2026 * d1, required_2026) == "A"


def test_boundary_d2_is_b(required_2026):
    d2 = CII_PARAMS["d_factors"]["d2"]
    assert cii.rate(required_2026 * d2, required_2026) == "B"


def test_boundary_d3_is_c(required_2026):
    d3 = CII_PARAMS["d_factors"]["d3"]
    assert cii.rate(required_2026 * d3, required_2026) == "C"


def test_boundary_d4_is_d(required_2026):
    d4 = CII_PARAMS["d_factors"]["d4"]
    assert cii.rate(required_2026 * d4, required_2026) == "D"


# ── compute_cii() 종합 (rating_ratio, boundaries 등) ────────────────────────
def test_compute_cii_success_matches_pure_functions():
    result = cii.compute_cii(scope="annual", year=2026, dwt=81200,
                              co2_mt=1000.0, distance_nm=10000.0)
    assert result.status == "success"
    assert approx(result.reference_cii, 4.192686272987268)
    assert approx(result.required_cii, 3.731490782958668)
    expected_attained = cii.attained_cii(1000.0, 81200, 10000.0)
    assert approx(result.attained_cii, expected_attained)
    assert result.rating == cii.rate(expected_attained, result.required_cii)
    assert approx(result.rating_ratio, expected_attained / result.required_cii)
    assert result.d1 == CII_PARAMS["d_factors"]["d1"]
    assert approx(result.boundary_a_b, result.required_cii * result.d1)
    assert result.unit == "gCO2/(DWT\u00b7nm)"


# ── 6. 잘못된 입력 → 계산 불가(unavailable), 등급 임의 부여 금지 ─────────────
def test_dwt_zero_is_unavailable():
    result = cii.compute_cii(scope="annual", year=2026, dwt=0,
                              co2_mt=100.0, distance_nm=1000.0)
    assert result.status == "unavailable"
    assert result.rating is None
    assert result.reason


def test_distance_zero_is_unavailable():
    result = cii.compute_cii(scope="annual", year=2026, dwt=81200,
                              co2_mt=100.0, distance_nm=0)
    assert result.status == "unavailable"
    assert result.rating is None


def test_distance_negative_is_unavailable():
    result = cii.compute_cii(scope="annual", year=2026, dwt=81200,
                              co2_mt=100.0, distance_nm=-5)
    assert result.status == "unavailable"
    assert result.rating is None


def test_unsupported_year_is_unavailable_no_extrapolation():
    result = cii.compute_cii(scope="annual", year=2031, dwt=81200,
                              co2_mt=100.0, distance_nm=1000.0)
    assert result.status == "unavailable"
    assert result.rating is None
    assert "2031" in result.reason
    result2 = cii.compute_cii(scope="annual", year=2022, dwt=81200,
                               co2_mt=100.0, distance_nm=1000.0)
    assert result2.status == "unavailable"


def test_co2_none_is_unavailable():
    result = cii.compute_cii(scope="annual", year=2026, dwt=81200,
                              co2_mt=None, distance_nm=1000.0)
    assert result.status == "unavailable"
    assert result.rating is None


def test_empty_aggregate_is_unavailable():
    result = cii.compute_cii(scope="annual", year=2026, dwt=81200,
                              co2_mt=None, distance_nm=None)
    assert result.status == "unavailable"
    assert result.rating is None


def test_co2_zero_with_valid_distance_is_success_zero():
    result = cii.compute_cii(scope="annual", year=2026, dwt=81200,
                              co2_mt=0.0, distance_nm=1000.0)
    assert result.status == "success"
    assert result.attained_cii == 0.0
    assert result.rating == "A"


# ── 7. vessel.dwt DB 동기화 (idempotent) ────────────────────────────────────
def test_db_dwt_sync_idempotent(tmp_path):
    from agent.db_schema import get_conn
    from scripts.load_hodata import sync_vessel_dwt, VESSEL_IMO

    db_path = tmp_path / "test_maritime.db"
    conn = get_conn(db_path)
    try:
        # 최초 동기화
        sync_vessel_dwt(conn)
        row = conn.execute("SELECT dwt, type FROM vessel WHERE imo=?", (VESSEL_IMO,)).fetchone()
        assert row is not None
        assert float(row["dwt"]) == 81200
        assert row["type"] == "Bulk Carrier"

        # 재실행(idempotent) — 여러 번 실행해도 동일 결과
        sync_vessel_dwt(conn)
        sync_vessel_dwt(conn)
        row2 = conn.execute("SELECT dwt FROM vessel WHERE imo=?", (VESSEL_IMO,)).fetchone()
        assert float(row2["dwt"]) == 81200
    finally:
        conn.close()


def test_db_dwt_never_reset_to_zero_by_reload(tmp_path):
    """load_hodata 재실행(rebuild 시나리오)에서도 vessel.dwt가 0으로 되돌아가지 않아야 한다."""
    from agent.db_schema import get_conn
    from scripts.load_hodata import sync_vessel_dwt, VESSEL_IMO, VESSEL_NAME

    db_path = tmp_path / "test_maritime2.db"
    conn = get_conn(db_path)
    try:
        # rebuild()가 초기화 시 수행하는 것과 동일한 시퀀스 재현
        conn.execute("DELETE FROM vessel")
        conn.execute(
            "INSERT OR IGNORE INTO vessel (imo, name, type) VALUES (?,?,?)",
            (VESSEL_IMO, VESSEL_NAME, VESSEL["type"]),
        )
        conn.commit()
        row0 = conn.execute("SELECT dwt FROM vessel WHERE imo=?", (VESSEL_IMO,)).fetchone()
        assert float(row0["dwt"]) == 0  # 스키마 기본값은 0 (동기화 전)

        sync_vessel_dwt(conn)
        row1 = conn.execute("SELECT dwt FROM vessel WHERE imo=?", (VESSEL_IMO,)).fetchone()
        assert float(row1["dwt"]) == 81200

        # dwt<=0을 명시적으로 넘기면 덮어쓰지 않고 건너뛴다 (0으로 되돌리지 않음)
        sync_vessel_dwt(conn, dwt=0)
        row2 = conn.execute("SELECT dwt FROM vessel WHERE imo=?", (VESSEL_IMO,)).fetchone()
        assert float(row2["dwt"]) == 81200
    finally:
        conn.close()


# ── 8. Word 보고서 통합 테스트 ───────────────────────────────────────────────
@pytest.fixture
def sample_cii_result():
    return cii.compute_cii(
        scope="annual", year=2026, dwt=81200, co2_mt=15000.0, distance_nm=45000.0,
    ).to_dict()


def _docx_text(path: Path) -> str:
    import docx
    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for c in row.cells:
                parts.append(c.text)
    return "\n".join(parts)


def test_mrv_annual_report_contains_cii(tmp_path, sample_cii_result, monkeypatch):
    from agent import reports

    monkeypatch.setattr(reports, "REPORTS_DIR", tmp_path)
    vessel = {"name": "H2521", "imo": "H2521", "type": "Bulk Carrier", "flag": "Unknown",
              "gt": 0, "dwt": 81200}
    summary = {
        "voyages_count": 3, "total_distance_nm": 45000.0, "total_days_at_sea": 120.0,
        "total_foc_oil_mt": 900.0, "total_fgc_gas_mt": 0.0,
        "total_co2_mt": 15000.0, "total_ch4_mt": 1.2, "total_co2e_mt": 15200.0,
        "total_cargo_mt": 0.0,
    }
    path = reports.generate_mrv_annual_docx(summary, [], vessel, 2026, sample_cii_result)
    assert path.exists()
    text = _docx_text(path)
    assert "Attained CII" in text
    assert "Required CII" in text
    assert "CII Rating" in text
    assert "81,200" in text
    assert sample_cii_result["rating"] in ("A", "B", "C", "D", "E")
    assert sample_cii_result["rating"] in text


def test_mrv_voyage_report_contains_indicative_cii(tmp_path, sample_cii_result, monkeypatch):
    from agent import reports

    monkeypatch.setattr(reports, "REPORTS_DIR", tmp_path)
    vessel = {"name": "H2521", "imo": "H2521", "type": "Bulk Carrier", "flag": "Unknown",
              "gt": 0, "dwt": 81200}
    voyage = {
        "voyage_id": "H2521_V01_Laden", "departure_port": "Unknown", "arrival_port": "Unknown",
        "departure_date": "2026-01-01", "arrival_date": "2026-01-20",
        "distance_nm": 5000.0, "days_at_sea": 19.0,
        "foc_oil_mt": 100.0, "fgc_gas_mt": 0.0,
        "co2_mt": 1500.0, "ch4_mt": 0.1, "co2e_mt": 1520.0,
    }
    voyage_cii = cii.compute_cii(scope="voyage", year=2026, dwt=81200,
                                  co2_mt=1500.0, distance_nm=5000.0).to_dict()
    path = reports.generate_mrv_voyage_docx(voyage, vessel, voyage_cii)
    assert path.exists()
    text = _docx_text(path)
    assert "Attained CII" in text
    assert "Required CII" in text
    assert "CII Rating" in text
    assert "Indicative Voyage CII" in text


def test_noon_report_contains_cii(tmp_path, sample_cii_result, monkeypatch):
    from agent import reports

    monkeypatch.setattr(reports, "REPORTS_DIR", tmp_path)
    vessel = {"name": "H2521", "imo": "H2521", "type": "Bulk Carrier", "flag": "Unknown",
              "gt": 0, "dwt": 81200}
    voyage = {"departure_port": "Unknown", "arrival_port": "Unknown", "voyage_id": "H2521_V01"}
    noon = {
        "report_datetime": "2026-06-28 12:00:00", "lat": 10.0, "lon": 100.0, "cog_deg": 90.0,
        "sog_kts": 12.0, "sailed_nm": 288.0, "me_rpm_note": "미측정", "me_power_kw": 5000.0,
        "foc_oil_mt": 10.0, "fgc_gas_mt": 0.0, "co2_mt": 30.0, "ch4_mt": 0.01, "co2e_mt": 31.0,
        "cii_value": 104.0, "wind_speed_kts": 0.0, "wind_dir_deg": 0.0, "wave_height_m": 0.0,
    }
    path = reports.generate_noon_report_docx(noon, voyage, vessel, sample_cii_result)
    assert path.exists()
    text = _docx_text(path)
    assert "CII (Carbon Intensity Indicator)" in text
    assert "Attained CII" in text


def test_report_unavailable_cii_does_not_crash(tmp_path, monkeypatch):
    from agent import reports

    monkeypatch.setattr(reports, "REPORTS_DIR", tmp_path)
    vessel = {"name": "H2521", "imo": "H2521", "type": "Bulk Carrier", "flag": "Unknown",
              "gt": 0, "dwt": 81200}
    unavailable = cii.compute_cii(scope="annual", year=2031, dwt=81200,
                                   co2_mt=1.0, distance_nm=1.0).to_dict()
    summary = {"voyages_count": 0, "total_distance_nm": 0, "total_days_at_sea": 0,
               "total_foc_oil_mt": 0, "total_fgc_gas_mt": 0, "total_co2_mt": 0,
               "total_ch4_mt": 0, "total_co2e_mt": 0, "total_cargo_mt": 0}
    path = reports.generate_mrv_annual_docx(summary, [], vessel, 2031, unavailable)
    assert path.exists()
    text = _docx_text(path)
    assert "Unavailable" in text
