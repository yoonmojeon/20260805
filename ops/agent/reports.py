"""
Word 문서 리포트 생성 모듈 (python-docx)
- Noon Report
- MRV Voyage Report
- MRV Annual Report
"""
from datetime import datetime, timezone
from pathlib import Path

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))
from config import REPORTS_DIR, VESSEL

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _header_row(table, texts: list, bg: str = "1E88E5"):
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = text
        _set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)


def _add_row(table, values: list, bold_first: bool = False):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(9)
        if bold_first and i == 0:
            run.bold = True
    return row


def _title(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)
    return h


def _kv_table(doc, pairs: list):
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(pairs):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = str(v)
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(row.cells[0], "E3F2FD")
    return table


def _fmt(val, decimals=2, unit=""):
    try:
        v = float(val or 0)
        return f"{v:.{decimals}f}{' ' + unit if unit else ''}"
    except Exception:
        return str(val or "-")


_CII_SCOPE_LABEL = {
    "annual":         "Annual (공식 연간 CII)",
    "current_voyage": "Current Voyage (잠정/Indicative)",
    "voyage":         "Voyage (참고/Indicative)",
}

_CII_RATING_COLOR = {
    "A": "1E88E5",  # 파랑
    "B": "43A047",  # 초록
    "C": "9E9E9E",  # 중립
    "D": "FB8C00",  # 주황
    "E": "E53935",  # 빨강
}


def _add_cii_section(doc, cii: dict, level: int = 2, title: str = "CII (Carbon Intensity Indicator)"):
    """
    공통 CII 섹션 렌더링. cii는 agent/cii.py compute_cii()의 결과(dict)이며,
    이 함수는 값을 표에 그대로 표시만 하고 재계산하지 않는다.
    """
    _title(doc, title, level=level)

    if not cii or cii.get("status") != "success":
        reason = (cii or {}).get("reason", "필요한 데이터가 없습니다.")
        _kv_table(doc, [
            ("CII Scope", _CII_SCOPE_LABEL.get((cii or {}).get("scope"), "-")),
            ("Status", "Unavailable"),
            ("Note", reason),
        ])
        return

    headers = ["Item", "Value"]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _header_row(table, headers)

    scope_label = _CII_SCOPE_LABEL.get(cii.get("scope"), cii.get("scope", "-"))
    rating = cii.get("rating", "-")
    ratio_pct = (cii.get("rating_ratio") or 0) * 100

    rows = [
        ("CII Scope",        scope_label),
        ("Reporting Year",   str(cii.get("year", "-"))),
        ("Ship Type",        cii.get("ship_type", "-")),
        ("DWT",              f"{float(cii.get('dwt', 0)):,.0f}"),
        ("CO2 Emissions",    _fmt(cii.get("co2_mt"), 2, "MT")),
        ("Distance Travelled", _fmt(cii.get("distance_nm"), 1, "nm")),
        ("Attained CII",     _fmt(cii.get("attained_cii"), 3, "gCO2/(DWT\u00b7nm)")),
        ("Reference CII",    _fmt(cii.get("reference_cii"), 3, "gCO2/(DWT\u00b7nm)")),
        ("Required CII",     _fmt(cii.get("required_cii"), 3, "gCO2/(DWT\u00b7nm)")),
        ("Reduction Factor", _fmt(cii.get("reduction_factor_percent"), 3, "%")),
        ("Rating Ratio",     _fmt(ratio_pct, 1, "%")),
        ("CII Rating",       rating),
        ("Unit",             cii.get("unit", "gCO2/(DWT\u00b7nm)")),
        ("Note",             cii.get("warning") or "-"),
    ]
    for k, v in rows:
        row = _add_row(table, [k, v])
        if k == "CII Rating":
            color_hex = _CII_RATING_COLOR.get(rating)
            if color_hex:
                _set_cell_bg(row.cells[1], color_hex)
                run = row.cells[1].paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if rating != "C" else RGBColor(0x00, 0x00, 0x00)


# ─────────────────────────────────────────────────────────────────────────────
# Noon Report
# ─────────────────────────────────────────────────────────────────────────────
def generate_noon_report_docx(data: dict, voyage: dict, vessel: dict, cii: dict = None) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.0)

    rdt = str(data.get("report_datetime", ""))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("NOON REPORT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"{vessel.get('name', VESSEL['name'])}  |  "
        f"IMO {vessel.get('imo', VESSEL['imo'])}  |  "
        f"{rdt[:10]}"
    )

    doc.add_paragraph()
    _title(doc, "1. 선박 정보", level=2)
    _kv_table(doc, [
        ("선박명",   vessel.get("name", VESSEL["name"])),
        ("IMO 번호", vessel.get("imo",  VESSEL["imo"])),
        ("선종",     vessel.get("type", "Container Ship")),
        ("선적국",   vessel.get("flag", "Unknown")),
        ("GT / DWT", f"{int(vessel.get('gt', 0)):,} / {int(vessel.get('dwt', 0)):,} MT"),
    ])

    doc.add_paragraph()
    _title(doc, "2. 위치 및 항로 정보", level=2)
    lat = float(data.get("lat", 0) or 0)
    lon = float(data.get("lon", 0) or 0)
    lat_str = f"{'N' if lat >= 0 else 'S'} {abs(lat):.4f}°"
    lon_str = f"{'E' if lon >= 0 else 'W'} {abs(lon):.4f}°"
    cog = float(data.get("cog_deg", 0) or 0)
    _kv_table(doc, [
        ("보고 일시 (UTC)",     rdt),
        ("위도 (Latitude)",     lat_str),
        ("경도 (Longitude)",    lon_str),
        ("침로/COG (Heading)",  f"{cog:.1f}°"),
        ("출발항",              voyage.get("departure_port", "-")),
        ("도착항",              voyage.get("arrival_port",   "-")),
        ("금일 항주거리 (Sailed Distance)", _fmt(data.get("sailed_nm",  0), 1, "nm")),
        ("배수량 (Displacement)", "미제공" if not data.get("displacement_mt") else _fmt(data.get("displacement_mt", 0), 0, "MT")),
        ("항차 ID",             voyage.get("voyage_id",     "-")),
    ])

    doc.add_paragraph()
    _title(doc, "3. 주기관 성능", level=2)
    _kv_table(doc, [
        ("M/E RPM",          data.get("me_rpm_note", _fmt(data.get("me_rpm", 0), 1, "rpm"))),
        ("M/E 출력",         _fmt(data.get("me_power_kw", 0), 0, "kW")),
        ("선속 (SOG)",       _fmt(data.get("sog_kts",     0), 1, "knots")),
        ("M/E FOC (Oil)",    _fmt(data.get("foc_oil_mt",  0), 3, "MT/day")),
        ("M/E FGC (Gas)",    _fmt(data.get("fgc_gas_mt",  0), 3, "MT/day")),
    ])

    doc.add_paragraph()
    _title(doc, "4. 배출량", level=2)
    _kv_table(doc, [
        ("CO2",              _fmt(data.get("co2_mt",  0), 2, "MT/day")),
        ("CH4",              _fmt(data.get("ch4_mt",  0), 4, "MT/day")),
        ("CO2e",             _fmt(data.get("co2e_mt", 0), 2, "MT/day")),
        ("CII (kg CO2/nm)",  _fmt(data.get("cii_value", 0), 4, "kg/nm")),
    ])

    doc.add_paragraph()
    _add_cii_section(doc, cii, level=2, title="5. CII (Carbon Intensity Indicator)")

    doc.add_paragraph()
    _title(doc, "6. 기상 및 해상 상태", level=2)
    wd = float(data.get("wind_dir_deg", 0) or 0)
    compass = ["N","NNE","NE","ENE","E","ESE","SE","SSE","S","SSW","SW","WSW","W","WNW","NW","NNW"]
    wind_dir_str = compass[int((wd + 11.25) / 22.5) % 16]
    _kv_table(doc, [
        ("풍속 (Wind Speed)",    _fmt(data.get("wind_speed_kts", 0), 1, "knots")),
        ("풍향 (Wind Direction)", f"{wind_dir_str} ({wd:.0f}°)"),
        ("파고 (Wave Height)",   _fmt(data.get("wave_height_m",  0), 2, "m")),
    ])

    doc.add_paragraph()
    doc.add_paragraph(
        "Draft — Master Confirmation Required   |   "
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
        style="Normal"
    ).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    imo_tag = vessel.get("imo", VESSEL["imo"])
    fname = f"NoonReport_{imo_tag}_{rdt[:10].replace('-','')}.docx"
    fpath = REPORTS_DIR / fname
    doc.save(str(fpath))
    return fpath


# ─────────────────────────────────────────────────────────────────────────────
# MRV Voyage Report
# ─────────────────────────────────────────────────────────────────────────────
def generate_mrv_voyage_docx(voyage: dict, vessel: dict, cii: dict = None) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MRV VOYAGE REPORT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"Regulation (EU) 2015/757  |  "
        f"{voyage.get('departure_port','')} → {voyage.get('arrival_port','')}  |  "
        f"Voyage ID: {voyage.get('voyage_id','')}"
    )

    doc.add_paragraph()
    _title(doc, "Part A. Ship & Voyage Identification", level=2)
    _kv_table(doc, [
        ("Ship Name",      vessel.get("name", VESSEL["name"])),
        ("IMO Number",     vessel.get("imo",  VESSEL["imo"])),
        ("Ship Type",      vessel.get("type", "Container Ship")),
        ("Flag State",     vessel.get("flag", "Unknown")),
        ("Gross Tonnage",  f"{int(vessel.get('gt', 0)):,} GT"),
        ("Voyage ID",      voyage.get("voyage_id", "")),
        ("Departure Port", voyage.get("departure_port", "")),
        ("Departure Date", str(voyage.get("departure_date", ""))[:10]),
        ("Arrival Port",   voyage.get("arrival_port", "")),
        ("Arrival Date",   str(voyage.get("arrival_date",   ""))[:10]),
    ])

    doc.add_paragraph()
    _title(doc, "Part B. Voyage Statistics", level=2)
    days = float(voyage.get("days_at_sea", 0) or 0)
    _kv_table(doc, [
        ("Distance Sailed",  f"{float(voyage.get('distance_nm', 0)):,.1f} nm"),
        ("Time at Sea",      f"{days:.0f} days ({days*24:.0f} hours)"),
    ])

    doc.add_paragraph()
    _title(doc, "Part C. Fuel Consumption & Emissions", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    _header_row(table, ["Fuel Type", "Consumed (MT)", "EF (tCO2/t)", "CO2 (MT)"])

    foc_oil = float(voyage.get("foc_oil_mt", 0) or 0)
    fgc_gas = float(voyage.get("fgc_gas_mt", 0) or 0)
    co2_tot = float(voyage.get("co2_mt",     0) or 0)
    ch4_mt  = float(voyage.get("ch4_mt",     0) or 0)
    co2e_mt = float(voyage.get("co2e_mt",    0) or 0)

    if foc_oil > 0:
        _add_row(table, ["Oil (VLSFO/LSMGO)", f"{foc_oil:.3f}", "3.114~3.206", f"{foc_oil*3.151:.2f}"])
    if fgc_gas > 0:
        _add_row(table, ["LNG (Gas)",         f"{fgc_gas:.3f}", "2.750", f"{fgc_gas*2.75:.2f}"])
    _add_row(table, ["TOTAL", f"{foc_oil+fgc_gas:.3f}", "-", f"{co2_tot:.2f}"])

    doc.add_paragraph()
    _kv_table(doc, [
        ("Total CO2 Emissions",  f"{co2_tot:.2f} MT"),
        ("Total CH4 Emissions",  f"{ch4_mt:.4f} MT"),
        ("Total CO2e Emissions", f"{co2e_mt:.2f} MT"),
    ])

    doc.add_paragraph()
    _add_cii_section(doc, cii, level=2, title="Part D. Indicative Voyage CII")

    doc.add_paragraph()
    _title(doc, "Part E. Cargo & Transport Work", level=2)
    dist_nm = float(voyage.get("distance_nm", 0) or 1)
    cargo   = float(voyage.get("cargo_mt",    0) or 0)
    gt_val  = float(vessel.get("gt", 0) or 0)
    _kv_table(doc, [
        ("Cargo Carried",   f"{cargo:,.1f} MT"),
        ("Transport Work (cargo·nm)", f"{cargo * dist_nm:,.0f} MT·nm"),
        ("Transport Work (GT·nm)",    f"{gt_val * dist_nm:,.0f} GT·nm"),
    ])

    doc.add_paragraph()
    doc.add_paragraph(
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}  "
        f"|  Pursuant to Regulation (EU) 2015/757",
        style="Normal"
    )

    imo_tag = vessel.get("imo", VESSEL["imo"])
    vid = voyage.get("voyage_id", "UNKNOWN")
    fname = f"MRV_Voyage_{vid}_{imo_tag}.docx"
    fpath = REPORTS_DIR / fname
    doc.save(str(fpath))
    return fpath


# ─────────────────────────────────────────────────────────────────────────────
# MRV Annual Report
# ─────────────────────────────────────────────────────────────────────────────
def generate_mrv_annual_docx(summary: dict, voyages_list,
                               vessel: dict, year: int, cii: dict = None) -> Path:
    import pandas as _pd
    voyages_df = (voyages_list if isinstance(voyages_list, _pd.DataFrame)
                  else (_pd.DataFrame(voyages_list) if voyages_list else _pd.DataFrame()))

    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"MRV ANNUAL REPORT  —  {year}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"Regulation (EU) 2015/757  |  "
        f"Reporting Period: {year}-01-01 ~ {year}-12-31"
    )

    doc.add_paragraph()
    _title(doc, "Part A. Ship Identification", level=2)
    _kv_table(doc, [
        ("Ship Name",  vessel.get("name", VESSEL["name"])),
        ("IMO Number", vessel.get("imo",  VESSEL["imo"])),
        ("Ship Type",  vessel.get("type", "Container Ship")),
        ("Flag State", vessel.get("flag", "Unknown")),
        ("GT",         f"{int(vessel.get('gt',  0)):,} GT"),
        ("DWT",        f"{int(vessel.get('dwt', 0)):,} MT"),
    ])

    doc.add_paragraph()
    _title(doc, "Part B. Annual Summary", level=2)
    foc_oil = summary.get("total_foc_oil_mt", 0)
    fgc_gas = summary.get("total_fgc_gas_mt", 0)
    _kv_table(doc, [
        ("Reporting Period",              f"{year}-01-01 ~ {year}-12-31"),
        ("Total Voyages",                 str(summary.get("voyages_count", 0))),
        ("Total Distance Travelled",      f"{summary.get('total_distance_nm', 0):,} nm"),
        ("Total Time at Sea",             f"{summary.get('total_days_at_sea', 0)} days"),
        ("Fuel Consumption — Oil (MT)",   f"{foc_oil:,.2f}"),
        ("Fuel Consumption — Gas (MT)",   f"{fgc_gas:,.2f}"),
        ("Total Fuel Consumed (MT)",      f"{foc_oil + fgc_gas:,.2f}"),
        ("Total CO2 Emissions (MT)",      f"{summary.get('total_co2_mt', 0):,.2f}"),
        ("Total CH4 Emissions (MT)",      f"{summary.get('total_ch4_mt', 0):,.4f}"),
        ("Total CO2e Emissions (MT)",     f"{summary.get('total_co2e_mt', 0):,.2f}"),
        ("Total Cargo Carried (MT)",      f"{summary.get('total_cargo_mt', 0):,.0f}"),
        ("Total Transport Work (GT·nm)",  f"{summary.get('total_transport_work', 0):,.0f}"),
    ])

    doc.add_paragraph()
    _add_cii_section(doc, cii, level=2, title="Part C. Annual CII Rating")

    if not voyages_df.empty:
        doc.add_paragraph()
        _title(doc, "Part D. Voyage Detail", level=2)
        headers = ["Voyage ID", "From", "To", "Departure", "Dist (nm)",
                   "Oil (MT)", "Gas (MT)", "CO2 (MT)", "CO2e (MT)"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        _header_row(table, headers)

        for _, row in voyages_df.iterrows():
            dep_port = row.get("departure_port", "-")
            arr_port = row.get("arrival_port",   "-")
            dep_date = str(row.get("departure_date", ""))[:10]
            dist     = float(row.get("distance_nm", 0) or 0)
            oil      = float(row.get("foc_oil_mt",  0) or 0)
            gas      = float(row.get("fgc_gas_mt",  0) or 0)
            co2      = float(row.get("co2_mt",      0) or 0)
            co2e     = float(row.get("co2e_mt",     0) or 0)
            _add_row(table, [
                row.get("voyage_id", ""),
                dep_port, arr_port, dep_date,
                f"{dist:,.0f}",
                f"{oil:.1f}", f"{gas:.1f}",
                f"{co2:.1f}", f"{co2e:.1f}",
            ])

    doc.add_paragraph()
    doc.add_paragraph(
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}  "
        f"|  Pursuant to Regulation (EU) 2015/757",
        style="Normal"
    )

    imo_tag = vessel.get("imo", VESSEL["imo"])
    fname = f"MRV_Annual_{year}_{imo_tag}.docx"
    fpath = REPORTS_DIR / fname
    doc.save(str(fpath))
    return fpath
